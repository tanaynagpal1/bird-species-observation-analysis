"""
Render the report to PDF and Word.

The content lives in report_content.build_blocks(). This module only knows how
to draw it. Keeping the two apart is what guarantees the PDF, the Word file and
the dashboard's Report page say exactly the same thing - a reader who found
them disagreeing would have no reason to trust any of them.

Run from the project root:

    python src/make_report.py

Writes:
    docs/Bird_Species_Observation_Analysis_Report.pdf
    docs/Bird_Species_Observation_Analysis_Report.docx
"""
from __future__ import annotations

import html
import io
import re
import sys
from datetime import date
from pathlib import Path

import pandas as pd

ROOT = Path(__file__).resolve().parent.parent
if str(ROOT / "src") not in sys.path:
    sys.path.insert(0, str(ROOT / "src"))
# app/theme.py is the dashboard's single source of colour. It is pure
# constants with no Streamlit import, so the PDF can read the real palette
# rather than keeping its own copy - which would drift the first time anyone
# retuned the dashboard.
if str(ROOT / "app") not in sys.path:
    sys.path.insert(0, str(ROOT / "app"))

import report_content as rc  # noqa: E402
import theme  # noqa: E402

from reportlab.lib import colors  # noqa: E402
from reportlab.lib.enums import TA_JUSTIFY  # noqa: E402
from reportlab.lib.pagesizes import A4  # noqa: E402
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet  # noqa: E402
from reportlab.lib.units import mm  # noqa: E402
from reportlab.pdfbase.pdfmetrics import stringWidth  # noqa: E402
from reportlab.platypus import (  # noqa: E402
    KeepTogether, PageBreak, Paragraph, SimpleDocTemplate, Spacer, Table,
    TableStyle,
)

# Every colour below is the dashboard's own, read from app/theme.py at import
# time. Retune the theme and the PDF follows automatically.
def _c(hex_string: str):
    return colors.HexColor(hex_string)


FOREST = _c(theme.FOREST)
GRASSLAND = _c(theme.GRASSLAND)
AT_RISK = _c(theme.AT_RISK)
INK = _c(theme.INK)
INK2 = _c(theme.INK2)
MUTED = _c(theme.MUTED)
RULE = _c(theme.BORDER)
BAND = _c(theme.PAGE)              # the dashboard's page tint, as row striping
CARD = _c(theme.CARD)
NOTE_BG = _c(theme.BANNER_BG)
NOTE_BORDER = _c(theme.BANNER_BORDER)
NOTE_EDGE = _c(theme.GRASSLAND)    # the banner's left rule on every page
NOTE_TEXT = _c(theme.BANNER_TEXT)
SIDEBAR_TOP = _c(theme.SIDEBAR_TOP)
SIDEBAR_BOTTOM = _c(theme.SIDEBAR_BOTTOM)
NAV_ACTIVE = _c(theme.NAV_ACTIVE)

VERDICT_COLOUR = {
    "finding": FOREST,
    "narrowed": GRASSLAND,
    "rejected": AT_RISK,
    "descriptive": MUTED,
    "method": NAV_ACTIVE,
}

PDF_OUT = ROOT / "docs" / "Bird_Species_Observation_Analysis_Report.pdf"
DOCX_OUT = ROOT / "docs" / "Bird_Species_Observation_Analysis_Report.docx"


# --------------------------------------------------------------- helpers
def _styles() -> dict:
    base = getSampleStyleSheet()
    s = {}
    s["body"] = ParagraphStyle(
        "body", parent=base["Normal"], fontName="Helvetica", fontSize=9.5,
        leading=14.5, textColor=INK, alignment=TA_JUSTIFY, spaceAfter=7,
    )
    s["h1"] = ParagraphStyle(
        "h1", parent=base["Heading1"], fontName="Helvetica-Bold", fontSize=16,
        leading=20, textColor=FOREST, spaceBefore=4, spaceAfter=10,
    )
    s["h2"] = ParagraphStyle(
        "h2", parent=base["Heading2"], fontName="Helvetica-Bold", fontSize=12,
        leading=16, textColor=INK, spaceBefore=12, spaceAfter=6,
    )
    s["h3"] = ParagraphStyle(
        "h3", parent=base["Heading3"], fontName="Helvetica-Bold", fontSize=10.5,
        leading=14, textColor=INK2, spaceBefore=9, spaceAfter=4,
    )
    s["bullet"] = ParagraphStyle(
        "bullet", parent=s["body"], leftIndent=13, bulletIndent=3, spaceAfter=5,
    )
    s["caption"] = ParagraphStyle(
        "caption", parent=base["Normal"], fontName="Helvetica-Oblique",
        fontSize=8, leading=11, textColor=MUTED, spaceBefore=3, spaceAfter=9,
    )
    s["cell"] = ParagraphStyle(
        "cell", parent=base["Normal"], fontName="Helvetica", fontSize=7.6,
        leading=10, textColor=INK,
    )
    s["cellhead"] = ParagraphStyle(
        "cellhead", parent=s["cell"], fontName="Helvetica-Bold",
        textColor=colors.white,
    )
    s["note"] = ParagraphStyle(
        "note", parent=s["body"], fontSize=9, leading=13.5,
        textColor=NOTE_TEXT, alignment=TA_JUSTIFY,
        spaceAfter=0,
    )
    s["verdict"] = ParagraphStyle(
        "verdict", parent=s["body"], fontSize=9.5, leading=14,
        textColor=INK, spaceAfter=0,
    )
    s["cover_title"] = ParagraphStyle(
        "ct", parent=base["Title"], fontName="Helvetica-Bold", fontSize=26,
        leading=31, textColor=FOREST, spaceAfter=10,
    )
    s["cover_sub"] = ParagraphStyle(
        "cs", parent=base["Normal"], fontName="Helvetica", fontSize=11.5,
        leading=17, textColor=INK2, spaceAfter=26,
    )
    s["cover_meta"] = ParagraphStyle(
        "cm", parent=base["Normal"], fontName="Helvetica", fontSize=9.5,
        leading=15, textColor=INK2,
    )
    return s


# Column names come straight from the analysis DataFrames, where snake_case is
# correct. In a printed report they need to read as English.
_COL_WORDS = {
    "pct": "%", "n": "n", "id": "ID", "rho": "rho", "tsn": "TSN",
    "pif": "PIF", "sd": "SD",
}
_COL_OVERRIDE = {
    "Admin_Unit_Code": "Park code",
    "Common_Name": "Species",
    "Scientific_Name": "Scientific name",
    "species_per_session": "Species / session",
    "sightings_per_session": "Sightings / session",
    "sessions_run": "Sessions run",
    "distinct_species": "Distinct species",
    "forest_higher": "Forest higher?",
    "both_habitats": "Both habitats?",
    "reliable": "Reliable (30+)?",
    "pct_of_all_at_risk": "% of all at-risk",
    "at_risk_sessions": "At-risk sessions",
    "pct_of_sessions": "% of sessions",
    "visits": "Visits",
    "total": "Total",
    "spread": "Spread",
    "mean": "Mean", "median": "Median", "min": "Min", "max": "Max",
    "count": "Count",
    "grassland_share_pct": "% in grassland",
    "total_sightings": "Sightings",
    "Plot_Name": "Plot",
    "time_band": "Time band",
    "month_name": "Month",
    "temp_band": "Temperature band",
    "humidity_band": "Humidity band",
    "session_duration_min": "Duration (min)",
}


def _pretty_col(name: str) -> str:
    """Turn a DataFrame column name into something a reader expects to see."""
    if name in _COL_OVERRIDE:
        return _COL_OVERRIDE[name]
    if not re.search(r"[a-z]_[a-z]|^[a-z]+$", name):
        return name          # already human-written, e.g. "Forest %"
    parts = name.split("_")
    out = [_COL_WORDS.get(p.lower(), p) for p in parts]
    text = " ".join(out)
    return text[:1].upper() + text[1:]


def _cols_of(df: pd.DataFrame) -> tuple[pd.DataFrame, list[str]]:
    """Reset the index when it carries meaning, and prettify the headers."""
    d = df.reset_index() if df.index.name or not isinstance(
        df.index, pd.RangeIndex) else df.copy()
    cols = []
    for c in d.columns:
        if isinstance(c, tuple):
            joined = " ".join(str(x) for x in c if str(x) != "").strip()
            cols.append(_pretty_col(joined))
        else:
            cols.append(_pretty_col(str(c)))
    return d, cols


def _esc(v) -> str:
    """Escape a raw value for a reportlab Paragraph."""
    if isinstance(v, float):
        v = f"{v:.4g}"
    return html.escape(str(v), quote=False)


def _df_cells(df: pd.DataFrame, st: dict) -> list[list]:
    """A DataFrame as Paragraph cells, so long text wraps instead of clipping."""
    d, cols = _cols_of(df)
    out = [[Paragraph(_esc(c), st["cellhead"]) for c in cols]]
    for _, row in d.iterrows():
        out.append([Paragraph(_esc(v), st["cell"]) for v in row])
    return out


def _table(df: pd.DataFrame, caption: str, st: dict, width: float):
    cells = _df_cells(df, st)
    ncols = len(cells[0])
    # Give the first column more room - it is nearly always the label column.
    if ncols > 2:
        first = width * 0.30
        rest = (width - first) / (ncols - 1)
        widths = [first] + [rest] * (ncols - 1)
    else:
        widths = [width / ncols] * ncols
    t = Table(cells, colWidths=widths, repeatRows=1, hAlign="LEFT")
    t.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), FOREST),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("TOPPADDING", (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
        ("LEFTPADDING", (0, 0), (-1, -1), 5),
        ("RIGHTPADDING", (0, 0), (-1, -1), 5),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [CARD, BAND]),
        ("GRID", (0, 0), (-1, -1), 0.4, RULE),
        ("LINEBELOW", (0, 0), (-1, 0), 0.8, FOREST),
    ]))
    return [t, Paragraph(_esc(caption), st["caption"])]


def _note(text: str, st: dict, width: float):
    t = Table([[Paragraph(text, st["note"])]], colWidths=[width], hAlign="LEFT")
    t.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, -1), NOTE_BG),
        ("BOX", (0, 0), (-1, -1), 0.5, NOTE_BORDER),
        ("LINEBEFORE", (0, 0), (0, -1), 2.5, NOTE_EDGE),
        ("TOPPADDING", (0, 0), (-1, -1), 8),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 8),
        ("LEFTPADDING", (0, 0), (-1, -1), 10),
        ("RIGHTPADDING", (0, 0), (-1, -1), 10),
    ]))
    return [Spacer(1, 3), t, Spacer(1, 9)]


def _verdict(kind: str, text: str, st: dict, width: float):
    colour = VERDICT_COLOUR.get(kind, MUTED)
    label = rc.VERDICT_LABELS.get(kind, kind.upper())
    chip = Table([[Paragraph(
        f'<font color="white" size="7"><b>{label}</b></font>', st["cell"])]],
        colWidths=[stringWidth(label, "Helvetica-Bold", 7) + 14],
        hAlign="LEFT")
    chip.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, -1), colour),
        ("TOPPADDING", (0, 0), (-1, -1), 3),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
        ("LEFTPADDING", (0, 0), (-1, -1), 6),
        ("RIGHTPADDING", (0, 0), (-1, -1), 6),
    ]))
    body = Table([[Paragraph(text, st["verdict"])]], colWidths=[width],
                 hAlign="LEFT")
    body.setStyle(TableStyle([
        ("LINEBEFORE", (0, 0), (0, -1), 2.5, colour),
        ("LEFTPADDING", (0, 0), (-1, -1), 10),
        ("RIGHTPADDING", (0, 0), (-1, -1), 0),
        ("TOPPADDING", (0, 0), (-1, -1), 3),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
    ]))
    return [Spacer(1, 2), chip, Spacer(1, 4), body, Spacer(1, 9)]


def _page_furniture(canvas, doc):
    """Header rule and page number on every page except the cover."""
    if doc.page == 1:
        _cover_band(canvas, doc)
    canvas.saveState()
    if doc.page > 1:
        canvas.setFont("Helvetica", 7.5)
        canvas.setFillColor(MUTED)
        canvas.drawString(20 * mm, A4[1] - 12 * mm, rc.TITLE)
        canvas.drawRightString(A4[0] - 20 * mm, A4[1] - 12 * mm,
                               date.today().strftime("%d %B %Y"))
        canvas.setStrokeColor(RULE)
        canvas.setLineWidth(0.5)
        canvas.line(20 * mm, A4[1] - 14 * mm, A4[0] - 20 * mm, A4[1] - 14 * mm)
        canvas.drawCentredString(A4[0] / 2, 12 * mm, str(doc.page - 1))
    canvas.restoreState()


def _cover_band(canvas, doc):
    """The dashboard's sidebar gradient, as a cover banner.

    Drawn on the canvas rather than in the story so it can bleed to the page
    edges the way the sidebar does on screen.
    """
    canvas.saveState()
    top, height = A4[1], 78 * mm
    steps = 220          # fine enough that no banding is visible in print
    r1, g1, b1 = SIDEBAR_TOP.red, SIDEBAR_TOP.green, SIDEBAR_TOP.blue
    r2, g2, b2 = SIDEBAR_BOTTOM.red, SIDEBAR_BOTTOM.green, SIDEBAR_BOTTOM.blue
    for i in range(steps):
        f = i / (steps - 1)
        canvas.setFillColorRGB(r1 + (r2 - r1) * f,
                               g1 + (g2 - g1) * f,
                               b1 + (b2 - b1) * f)
        canvas.rect(0, top - height + (height / steps) * i,
                    A4[0], height / steps + 1.2, stroke=0, fill=1)
    canvas.setFillColor(colors.white)
    canvas.setFont("Helvetica-Bold", 25)
    canvas.drawString(20 * mm, top - 40 * mm, rc.TITLE)
    canvas.setFont("Helvetica", 10.5)
    canvas.setFillColorRGB(0.81, 0.91, 0.85)
    text = canvas.beginText(20 * mm, top - 50 * mm)
    text.setLeading(15)
    for line in ("Forest and grassland point-count surveys",
                 "11 National Park Service units, 2018 breeding season"):
        text.textLine(line)
    canvas.drawText(text)
    canvas.restoreState()


def _cover(st: dict, ctx: dict, width: float) -> list:
    rows_, sessions_ = ctx["rows"], ctx["sessions"]
    q12 = ctx["r"]["q12"]
    # The title itself is painted by _cover_band on the canvas; the story
    # starts below the band.
    story = [Spacer(1, 68 * mm)]
    meta = [
        ["Dataset", f"{len(rows_):,} sightings, {len(sessions_):,} survey "
                    f"sessions, {rows_['Scientific_Name'].nunique()} species"],
        ["Coverage", f"{q12['n_parks']} National Park Service units, "
                     f"May to July 2018"],
        ["Habitats", "Forest and grassland point-count surveys"],
        ["Report generated", date.today().strftime("%d %B %Y")],
        ["Source", "analysis.run_all() - every figure computed at build time"],
    ]
    t = Table([[Paragraph(f"<b>{k}</b>", st["cover_meta"]),
                Paragraph(v, st["cover_meta"])] for k, v in meta],
              colWidths=[width * 0.26, width * 0.74], hAlign="LEFT")
    t.setStyle(TableStyle([
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("TOPPADDING", (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
        ("LEFTPADDING", (0, 0), (-1, -1), 0),
        ("LINEABOVE", (0, 0), (-1, 0), 0.8, FOREST),
        ("LINEBELOW", (0, -1), (-1, -1), 0.8, FOREST),
    ]))
    story += [t, PageBreak()]
    return story


def _contents(blocks: list, st: dict, width: float) -> list:
    """A table of contents built from the h1/h2 blocks themselves."""
    entries = [(b[0], b[1]) for b in blocks if b[0] in ("h1", "h2")]
    rowsl = []
    for kind, text in entries:
        indent = "" if kind == "h1" else "&nbsp;&nbsp;&nbsp;&nbsp;"
        weight = ("<b>%s</b>" if kind == "h1" else "%s") % _esc(text)
        rowsl.append([Paragraph(indent + weight, st["cell"])])
    t = Table(rowsl, colWidths=[width], hAlign="LEFT")
    t.setStyle(TableStyle([
        ("TOPPADDING", (0, 0), (-1, -1), 2.5),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 2.5),
        ("LEFTPADDING", (0, 0), (-1, -1), 0),
    ]))
    return [Paragraph("Contents", st["h1"]), t, PageBreak()]


def build_pdf(path: Path | None = None) -> bytes:
    """Render the report to PDF and return the bytes."""
    blocks, ctx = rc.build_blocks()
    st = _styles()
    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=A4,
        leftMargin=20 * mm, rightMargin=20 * mm,
        topMargin=20 * mm, bottomMargin=18 * mm,
        title=rc.TITLE, author="Bird Species Observation Analysis",
        subject=rc.SUBTITLE,
    )
    width = doc.width
    story = _cover(st, ctx, width) + _contents(blocks, st, width)

    for b in blocks:
        kind = b[0]
        if kind == "pagebreak":
            story.append(PageBreak())
        elif kind in ("h1", "h2", "h3"):
            story.append(Paragraph(_esc(b[1]), st[kind]))
        elif kind == "p":
            story.append(Paragraph(b[1], st["body"]))
        elif kind == "bullets":
            for item in b[1]:
                story.append(Paragraph(item, st["bullet"], bulletText="•"))
            story.append(Spacer(1, 4))
        elif kind == "kv":
            cells = [[Paragraph(f"<b>{_esc(k)}</b>", st["cell"]),
                      Paragraph(_esc(v), st["cell"])] for k, v in b[1]]
            t = Table(cells, colWidths=[width * 0.30, width * 0.70],
                      hAlign="LEFT")
            t.setStyle(TableStyle([
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("TOPPADDING", (0, 0), (-1, -1), 3),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
                ("LEFTPADDING", (0, 0), (-1, -1), 0),
                ("LINEBELOW", (0, 0), (-1, -2), 0.3, RULE),
            ]))
            story += [t, Spacer(1, 9)]
        elif kind == "table":
            story += _table(b[1], b[2], st, width)
        elif kind == "note":
            story += _note(b[1], st, width)
        elif kind == "verdict":
            story += _verdict(b[1], b[2], st, width)

    doc.build(story, onFirstPage=_page_furniture, onLaterPages=_page_furniture)
    data = buf.getvalue()
    if path:
        path.parent.mkdir(parents=True, exist_ok=True)
        path.write_bytes(data)
    return data


# --------------------------------------------------------------- word
_TAG = re.compile(r"(<b>|</b>|<i>|</i>)")


def _rich(par, text: str) -> None:
    """Write text into a docx paragraph, honouring <b> and <i>."""
    bold = italic = False
    for piece in _TAG.split(html.unescape(text)):
        if piece == "<b>":
            bold = True
        elif piece == "</b>":
            bold = False
        elif piece == "<i>":
            italic = True
        elif piece == "</i>":
            italic = False
        elif piece:
            run = par.add_run(piece)
            run.bold, run.italic = bold, italic


def build_docx(path: Path | None = None) -> bytes:
    """Render the report to a Word document and return the bytes."""
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, RGBColor

    blocks, ctx = rc.build_blocks()
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10)

    title = doc.add_heading(rc.TITLE, level=0)
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x12, 0x80, 0x5C)
    sub = doc.add_paragraph()
    sub.add_run(rc.SUBTITLE).italic = True
    meta = doc.add_paragraph()
    meta.add_run(
        f"{len(ctx['rows']):,} sightings across {len(ctx['sessions']):,} "
        f"survey sessions | generated "
        f"{date.today().strftime('%d %B %Y')}"
    ).font.size = Pt(9)
    doc.add_page_break()

    for b in blocks:
        kind = b[0]
        if kind == "pagebreak":
            doc.add_page_break()
        elif kind == "h1":
            doc.add_heading(b[1], level=1)
        elif kind == "h2":
            doc.add_heading(b[1], level=2)
        elif kind == "h3":
            doc.add_heading(b[1], level=3)
        elif kind == "p":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            _rich(p, b[1])
        elif kind == "bullets":
            for item in b[1]:
                _rich(doc.add_paragraph(style="List Bullet"), item)
        elif kind == "kv":
            t = doc.add_table(rows=0, cols=2)
            t.style = "Light List Accent 1"
            for k, v in b[1]:
                cells = t.add_row().cells
                cells[0].text = str(k)
                cells[1].text = str(v)
        elif kind == "table":
            d, cols = _cols_of(b[1])
            t = doc.add_table(rows=1, cols=len(cols))
            t.style = "Light Grid Accent 1"
            for i, c in enumerate(cols):
                t.rows[0].cells[i].text = str(c)
            for _, row in d.iterrows():
                cells = t.add_row().cells
                for i, v in enumerate(row):
                    cells[i].text = (f"{v:.4g}" if isinstance(v, float)
                                     else str(v))
            cap = doc.add_paragraph()
            cap_run = cap.add_run(b[2])
            cap_run.italic = True
            cap_run.font.size = Pt(8)
        elif kind == "note":
            p = doc.add_paragraph()
            _rich(p, b[1])
            for run in p.runs:
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0x6B, 0x53, 0x14)
        elif kind == "verdict":
            p = doc.add_paragraph()
            lab = p.add_run(rc.VERDICT_LABELS.get(b[1], b[1].upper()) + "  ")
            lab.bold = True
            lab.font.size = Pt(8)
            _rich(p, b[2])

    buf = io.BytesIO()
    doc.save(buf)
    data = buf.getvalue()
    if path:
        path.parent.mkdir(parents=True, exist_ok=True)
        path.write_bytes(data)
    return data


if __name__ == "__main__":
    build_pdf(PDF_OUT)
    print(f"wrote {PDF_OUT.relative_to(ROOT)}")
    build_docx(DOCX_OUT)
    print(f"wrote {DOCX_OUT.relative_to(ROOT)}")